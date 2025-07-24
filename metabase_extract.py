import os
import time
import re
import subprocess
import shutil

from PIL import Image
import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from docx import Document as docx_document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


class MetabaseDashboardExtract:
    def __init__(self, email, password, base_url, output_dir="output", model=None):
        self.email = email
        self.password = password
        self.base_url = base_url
        self.output_dir = output_dir
        self.model = model
        
        chrome_options = Options()
        chrome_options.add_argument("--headless=new")
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument("--window-size=1920,1080")
        chrome_options.add_experimental_option("prefs", {
            "profile.default_content_setting_values.notifications": 2,
            "profile.default_content_settings.popups": 0
        })

        self.driver = webdriver.Chrome(options=chrome_options)
        os.makedirs(self.output_dir, exist_ok=True)

    def login(self):
        login_url = f"{self.base_url}/auth/login"
        self.driver.get(login_url)
        try:
            username_field = WebDriverWait(self.driver, 20).until(EC.element_to_be_clickable((By.NAME, "username")))
            password_field = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.NAME, "password")))
            username_field.send_keys(self.email)
            password_field.send_keys(self.password)
            login_button = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "button[type='submit']")))
            login_button.click()
            WebDriverWait(self.driver, 30).until(EC.url_contains("/"))
        except TimeoutException as e:
            print(f"Error en login: {e}")
            raise

    def capture_dashboard(self, dashboard_url, municipio, tab):
        print(f"capturando dashboard para {municipio} - {tab}")
        self.driver.get(dashboard_url)
        WebDriverWait(self.driver, 30).until(EC.presence_of_element_located((By.CSS_SELECTOR, '[data-testid="dashcard-container"]')))
        time.sleep(5)

        cards = self.driver.find_elements(By.CSS_SELECTOR, "div.elbzci2t1.emotion-7apf2f1.e1isodme0") + \
                self.driver.find_elements(By.CSS_SELECTOR, '[data-testid="visualization-root"]')

        if not cards:
            print("No se encontraron gráficos.")
            return

        folder_path = os.path.join(self.output_dir, municipio, tab)
        os.makedirs(folder_path, exist_ok=True)

        for i, card in enumerate(cards):
            try:
                self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", card)
                time.sleep(1)
                WebDriverWait(self.driver, 10).until(EC.visibility_of(card))

                title = ""
                #SIMPLIFICAR
                for selector in ["[data-testid='legend-caption-title']", "[data-testid='scalar-title']", ".Card-title", "h3", "h4"]:
                    try:
                        title = card.find_element(By.CSS_SELECTOR, selector).text.strip()
                        if title:
                            break
                    except:
                        continue

                if not title:
                    title = f"grafico_{i+1}"

                clean_title = re.sub(r'[<>:"/\\|?*\n\r\t]', "_", title)[:80]
                filepath = os.path.join(folder_path, f"{clean_title}.png")
                card.screenshot(filepath)
                print(f"Capturado: {filepath}")
            except Exception as e:
                print(f"Error capturando gráfico {i+1}: {e}")
                
    
    
    def get_image_description_from_gemini(self, image_paths, prompt):
        pil_images = []
        try:
            for path in image_paths:
                img = Image.open(path)
                pil_images.append(img)
            response = self.model.generate_content([prompt] + pil_images)
            return getattr(response, 'text', None) or getattr(response.candidates[0], 'text', None)
        except Exception as e:
            print(f"Error al cargar las imágenes: {e}")
            return None
        
        finally:
            for i in pil_images:
                if i:
                    i.close()

    def select_relevant_images(self, image_folder, question):
        output_folder = os.path.join(image_folder, "seleccionadas")
        os.makedirs(output_folder, exist_ok=True)
        relevantes = []

        for file in os.listdir(image_folder):
            if file.endswith(".png"):
                path = os.path.join(image_folder, file)
                prompt = (
                    f"La siguiente imagen representa un gráfico extraído de datos turísticos. "
                    f"El usuario hizo esta pregunta: \"{question}\".\n"
                    f"¿Esta visualización es útil o relevante para responder a esa consulta? "
                    f"Respondé exclusivamente con una sola palabra: True o False. "
                )
                result = self.get_image_description_from_gemini([path], prompt)
                if result and "true" in result.lower():
                    shutil.copy(path, output_folder)
                    relevantes.append(path)

        return relevantes
    
    def export_to_docx(self, question, logo_path=None):
        doc_dir = os.path.join(self.output_dir, "docs")
        os.makedirs(doc_dir, exist_ok=True)

        for municipio in os.listdir(self.output_dir):
            municipio_path = os.path.join(self.output_dir, municipio)
            if not os.path.isdir(municipio_path):
                continue

            doc = docx_document()
            doc.styles['Normal'].font.name = 'Calibri'
            doc.styles['Normal'].font.size = Pt(12)

            
            title = doc.add_paragraph()
            run = title.add_run(f"Informe de {municipio.capitalize()}\n")
            run.font.size = Pt(32)
            run.bold = True
            run.underline = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            


            
            preambulo = (
                "Este informe presenta información turística específica del municipio, con el objetivo de responder a la siguiente pregunta clave:\n"
                f"{question}"
                "\nPara ello, se han seleccionado y analizado visualizaciones relevantes que permiten identificar tendencias y patrones relacionados. "
                "El análisis busca proporcionar evidencia clara y útil para la toma de decisiones estratégicas en el ámbito turístico."
            )
        
            preambulo_para = doc.add_paragraph()
            run = preambulo_para.add_run(preambulo)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            preambulo_para.paragraph_format.space_after = Pt(6)

            if logo_path and os.path.exists(logo_path):
                footer = doc.sections[0].footer
                table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
                table.autofit = False
                table.columns[0].width = Inches(3.25)
                table.columns[1].width = Inches(3.25)

                logo_cell = table.cell(0, 0)
                logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                logo_cell.paragraphs[0].add_run().add_picture(logo_path, width=Inches(0.75))
            
            for tab in sorted(os.listdir(municipio_path)):
                tab_path = os.path.join(municipio_path, tab)
                if not os.path.isdir(tab_path):
                    continue
                
                section_title = tab.replace("-", " ").replace("100", "").replace("102", "").replace("103", "").strip().title()
                if "vuts" in section_title.lower():
                    section_title = "Casas rurales y viviendas de uso turístico"
                elif "hoteles" in section_title.lower():
                    section_title = "Hoteles, hostales y campings"
                elif "oficiales" in section_title.lower():
                    section_title = "Datos de fuentes oficiales"

                section = doc.add_paragraph()
                section_run = section.add_run(section_title) 
                section_run.bold = True
                section.alignment = WD_ALIGN_PARAGRAPH.LEFT
                section.paragraph_format.space_after = Pt(16)
                

                selected_imgs = self.select_relevant_images(tab_path, question)
                for img_path in selected_imgs:
                        file = os.path.basename(img_path)

                        graph_title = doc.add_paragraph()
                        graph_title_run = graph_title.add_run(f"{os.path.splitext(file)[0].replace('_', ' ').title()}")
                        graph_title_run.font.name = 'Calibri'
                        graph_title_run.font.size = Pt(13)
                        graph_title_run.bold = True
                        graph_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        graph_title.paragraph_format.space_after = Pt(6)
                        

                        img_para = doc.add_paragraph()
                        img_run = img_para.add_run()
                        img_run.add_picture(img_path, width=Inches(6.0))
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.paragraph_format.space_after = Pt(12)
                        
                        
                        
                        prompt = (
                            f"Esta imagen será incluida en un informe de datos turísticos que busca responder la siguiente pregunta:\n"
                            f"\"{question}\"\n\n"
                            "Redactá una descripción clara, profesional y enfocada en cómo este gráfico aporta información relevante para responder esa pregunta. "
                            "No uses listas ni formato especial. Usá lenguaje técnico accesible, en formato de párrafo completo."
                        )
                            
                        description = self.get_image_description_from_gemini([img_path], prompt)

                        if description:
                            desc_para = doc.add_paragraph()
                            desc_run = desc_para.add_run(description.strip())
                            desc_run.font.size = Pt(12)
                            desc_run.font.name = 'Calibri'
                            desc_para.paragraph_format.space_after = Pt(6)

                        else:
                            doc.add_paragraph("(Fallo la descripción)")

                        
        conclusion_prompt = (
            f"Basándote en las visualizaciones y datos disponibles, redactá una conclusión profesional y clara que responda a la siguiente pregunta:\n"
            f"\"{question}\"\n\n"
            "Usá un tono técnico accesible, sin listas ni formato especial. Que parezca una conclusión escrita por un analista turístico."
        )

        
        conclusion_text = self.get_image_description_from_gemini([], conclusion_prompt)
        if conclusion_text:
            conclusion_heading = doc.add_paragraph("Conclusión", style='Heading 1')
            conclusion_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            conclusion_heading.paragraph_format.space_after = Pt(12)
            conclusion_para = doc.add_paragraph(conclusion_text.strip())
            conclusion_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        else:
            doc.add_paragraph("No se pudo generar una conclusión automática.")



        safe_municipio = municipio.replace(" ", "_").lower()
        output_path = os.path.join(doc_dir, f"{safe_municipio}.docx")
        doc.save(output_path)

    def docx_to_pdf(self):
        input_path = "output/docs/teulada.docx"
        output_dir = "output/pdf"
        os.makedirs(output_dir, exist_ok=True)

        try:
            subprocess.run([
                "D:/Programas/LibreOffice/program/soffice.exe",       
                "--headless",           
                "--convert-to", "pdf",  
                "--outdir", output_dir,
                input_path
            ], check=True)

            print("todo ok")
        except subprocess.CalledProcessError as e:
            print(f"error libreoffice: {e}")

    
    def run(self):
        inicio = time.time()
        try:
            self.login()
            #municipios = ["teulada", "benissa", "calpe", "javea", "mijas", "moraira", "platja_de_aro"]
            municipios = ["teulada"]
            #tabs = ["100-vuts-y-casas-rurales", "102-hoteles%2C-hostales-y-campings", "103-datos-de-fuentes-oficiales"]
            tabs = ["100-vuts-y-casas-rurales"]
            question = "¿Que tendencia se observa en la antelacion con la que los turistas reservan su alojamiento?"
            logo_path = "logo_municipio.png"

            for municipio in municipios:
                for tab in tabs:
                    dashboard_url = f"https://analytics.peninsula.co/dashboard/32-{municipio}?a%25C3%25B1o=&fecha=&mes=&per%25C3%25ADodo=&poblaci%25C3%25B3n=teulada&poblaci%25C3%25B3n_%28consumo%29=calpe&poblaci%25C3%25B3n_2=benissa&tab={tab}&tipo_de_alojamiento=&tipo_de_establecimiento="
                    try:
                        self.capture_dashboard(dashboard_url, municipio, tab)
                    except Exception as e:
                        print(f"ERROR: {municipio} - {tab}: {e}")
                        break
            
            self.export_to_docx(question, logo_path)
            self.docx_to_pdf()
            pdf_path = os.path.abspath("output/pdf/teulada.pdf")
            os.startfile(pdf_path)

        except Exception as e:
            print(f"Error general: {e}")
        finally:
            self.driver.quit()
            fin = time.time()
            print(fin-inicio)
        
        
"""if __name__ == "__main__":
    email = os.getenv("METABASE_EMAIL")
    password = os.getenv("METABASE_PASSWORD")
    base_url = os.getenv("METABASE_BASE_URL")
    api_key = os.getenv("GEMINI_API_KEY")

    if not email or not password or not api_key or not base_url:
        raise ValueError("No estan las variables de entorno en el sistema")
    
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.0-flash')

    MetabaseDashboardExtract(email, password, base_url).run()"""
    