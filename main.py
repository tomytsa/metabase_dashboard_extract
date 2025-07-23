import os
import time
import re
import subprocess
import shutil

from PIL import Image
import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from docx import Document as docx_document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


class MetabaseDashboardExtract:
    def __init__(self, email, password, base_url, output_dir="output"):
        self.email = email
        self.password = password
        self.base_url = base_url
        self.output_dir = output_dir

        
        chrome_options = Options()
        
        chrome_options.add_argument("--headless=new")
        
        chrome_options.add_argument("--disable-gpu")
        
        chrome_options.add_argument("--window-size=1920,1080")

        prefs = {
            "profile.default_content_setting_values.notifications": 2,
            "profile.default_content_settings.popups": 0
        }
        chrome_options.add_experimental_option("prefs", prefs)

        self.driver = webdriver.Chrome(options=chrome_options)
        os.makedirs(self.output_dir, exist_ok=True)

    def login(self):
        login_url = f"{self.base_url}/auth/login"
        self.driver.get(login_url)

        try:
            username_field = WebDriverWait(self.driver, 20).until(
                EC.element_to_be_clickable((By.NAME, "username"))
            )
            password_field = WebDriverWait(self.driver, 10).until(
                EC.element_to_be_clickable((By.NAME, "password"))
            )
            username_field.send_keys(self.email)
            password_field.send_keys(self.password)

            login_button = WebDriverWait(self.driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "button[type='submit']"))
            )
            login_button.click()

            WebDriverWait(self.driver, 30).until(
                EC.url_contains("/")
            )
        except TimeoutException as e:
            print(f"Error en login: {e}")
            raise

    def capture_dashboard(self, dashboard_url, municipio, tab):

        print(f"capturando dashboard para {municipio} - {tab}")

        self.driver.get(dashboard_url)
        WebDriverWait(self.driver, 30).until(EC.presence_of_element_located((By.CSS_SELECTOR, '[data-testid="dashcard-container"]')))
        time.sleep(10)

        
        cards = self.driver.find_elements(By.CSS_SELECTOR, "div.elbzci2t1.emotion-7apf2f1.e1isodme0")
        viz_cards = self.driver.find_elements(By.CSS_SELECTOR, '[data-testid="visualization-root"]')
        all_cards = cards + viz_cards

        if not all_cards:
            print("No se encontraron gráficos.")
            return


        for i, card in enumerate(all_cards):
            try:
                self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", card)
                time.sleep(2)
                WebDriverWait(self.driver, 10).until(EC.visibility_of(card))

                title = ""
                for selector in ["[data-testid='legend-caption-title']", "[data-testid='scalar-title']", ".Card-title", "h3", "h4"]:
                    try:
                        title_element = card.find_element(By.CSS_SELECTOR, selector)
                        title = title_element.text.strip()
                        if title:
                            break
                    except:
                        continue

                if not title:
                    title = f"grafico_{i+1}"

                clean_title = re.sub(r'[<>:"/\\|?*\n\r\t]', "_", title)[:80]
                folder_path = os.path.join(self.output_dir, municipio, tab)
                os.makedirs(folder_path, exist_ok=True)

                filename = f"{clean_title}.png"
                filepath = os.path.join(folder_path, filename)

                card.screenshot(filepath)
                print(f"Capturado: {filepath}")

            except Exception as e:
                print(f"Error capturando gráfico {i+1}: {e}")
                continue
    
    @staticmethod
    def get_image_description_from_gemini(image_paths, prompt):
        pil_images = []
        try:
            for path in image_paths:
                img = Image.open(path)
                pil_images.append(img)
        except Exception as e:
            print(f"Error al cargar las imágenes: {e}")
            return None
        
        try:
            response = model.generate_content([prompt] + pil_images)
        
            if hasattr(response, 'text'):
                return response.text
            elif response and response.candidates and response.candidates[0] and hasattr(response.candidates[0], 'text'):
                return response.candidates[0].text
            else:
                print(f"Gemini no devolvió un texto válido para las imágenes.")
                return None
        except Exception as e:
            print(f"Error al procesar las imágenes {', '.join([os.path.basename(p) for p in image_paths])} con Gemini: {e}")
            return None
        finally:
            for img in pil_images:
                if img:
                    img.close()

    def select_relevant_images(self, image_folder, question, output_folder=None):
        relevant_images = []
        all_files = os.listdir(image_folder)
        all_images = []

        for file_name in all_files:
            if file_name.lower().endswith(".png"):
                all_images.append(file_name)

        if not output_folder:
            output_folder = os.path.join(image_folder, "seleccionadas")

        os.makedirs(output_folder, exist_ok=True)
        
        for image_file in all_images:
            image_path = os.path.join(image_folder, image_file)

            prompt = (
                f"La siguiente imagen representa un gráfico extraído de datos turísticos. "
                f"El usuario hizo esta pregunta: \"{question}\".\n"
                f"¿Esta visualización es útil o relevante para responder a esa consulta? "
                f"Respondé exclusivamente con una sola palabra: True o False. "
            )

            try:
                response = self.get_image_description_from_gemini([image_path], prompt)
                print(f"{image_file}: {response}")
                if response and "true" in response.lower():
                    shutil.copy(image_path, output_folder)
                    relevant_images.append(image_path)
            except Exception as e:
                print(f"Error evaluando relevancia de {image_file}: {e}")

        return relevant_images
    
    def export_to_docx(self, question, logo_path=None):
        
        doc_dir = os.path.join(self.output_dir, "docs")
        os.makedirs(doc_dir, exist_ok=True)

        municipios = os.listdir(self.output_dir)
        for municipio in municipios:
            municipio_path = os.path.join(self.output_dir, municipio)
            if not os.path.isdir(municipio_path):
                continue

            doc = docx_document()

            doc.styles['Normal'].font.name = 'Calibri'
            doc.styles['Normal'].font.size = Pt(12)

            # Titulo
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"Informe de {municipio.capitalize()}\n")
            title_run.font.name = 'Calibri'
            title_run.font.size = Pt(32)
            title_run.bold = True
            title_run.underline = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            


            # Preambulo
            preambulo_text = (
                "Este informe presenta información turística específica del municipio, con el objetivo de responder a la siguiente pregunta clave:\n"
                f"{question}"
                "\nPara ello, se han seleccionado y analizado visualizaciones relevantes que permiten identificar tendencias y patrones relacionados. "
                "El análisis busca proporcionar evidencia clara y útil para la toma de decisiones estratégicas en el ámbito turístico."
            )
            preambulo_para = doc.add_paragraph()
            preambulo_run = preambulo_para.add_run(preambulo_text)
            preambulo_run.font.name = 'Calibri'
            preambulo_run.font.size = Pt(12)
            preambulo_para.paragraph_format.space_after = Pt(6)


            ## Footer para la primera sección
            section = doc.sections[0]
            section.footer.is_linked_to_previous = False 

            footer = section.footer
            footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
            footer_table.autofit = False
            footer_table.columns[0].width = Inches(3.25)
            footer_table.columns[1].width = Inches(3.25)

            # Logo (columna izquierda)
            logo_cell = footer_table.cell(0, 0)
            logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if logo_path and os.path.exists(logo_path):
                try:
                    logo_run = logo_cell.paragraphs[0].add_run()
                    logo_run.add_picture(logo_path, width=Inches(0.75))
                    logo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                except Exception as e:
                    print(f"Error al insertar el logo en el footer: {e}")

            

            # Páginas siguientes por cada tab / sección
            tabs_list = sorted(os.listdir(municipio_path))

            for tab_index, tab in enumerate(tabs_list):
                tab_path = os.path.join(municipio_path, tab)
                if not os.path.isdir(tab_path):
                    continue
            
                
                section_title_text = tab.replace("-", " ").replace("100", "").replace("102", "").replace("103", "").strip().title()
                if "vuts" in section_title_text.lower():
                    section_title_text = "Casas rurales y viviendas de uso turístico"
                elif "hoteles" in section_title_text.lower():
                    section_title_text = "Hoteles, hostales y campings"
                elif "oficiales" in section_title_text.lower():
                    section_title_text = "Datos de fuentes oficiales"

                section_heading = doc.add_paragraph()
                section_heading_run = section_heading.add_run(section_title_text) 
                section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                section_heading.paragraph_format.space_after = Pt(16)
                section_heading_run.bold = True

                selected_imgs = self.select_relevant_images(tab_path, question)
                for img_path in selected_imgs:
                        file = os.path.basename(img_path)


                        graph_title_para = doc.add_paragraph()
                        graph_title_run = graph_title_para.add_run(f"{os.path.splitext(file)[0].replace('_', ' ').title()}")
                        graph_title_run.font.name = 'Calibri'
                        graph_title_run.font.size = Pt(13)
                        graph_title_run.bold = True
                        graph_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        graph_title_para.paragraph_format.space_after = Pt(6)
                        

                        # Image
                        img_para = doc.add_paragraph()
                        img_run = img_para.add_run()
                        img_run.add_picture(img_path, width=Inches(6.0))
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.paragraph_format.space_after = Pt(12)
                        
                        
                        try:
                           
                            prompt = (
                                f"Esta imagen será incluida en un informe de datos turísticos que busca responder la siguiente pregunta:\n"
                                f"\"{question}\"\n\n"
                                "Redactá una descripción clara, profesional y enfocada en cómo este gráfico aporta información relevante para responder esa pregunta. "
                                "No uses listas ni formato especial. Usá lenguaje técnico accesible, en formato de párrafo completo."
                                )
                            description = MetabaseDashboardExtract.get_image_description_from_gemini([img_path], prompt)
                            time.sleep(0.5)

                            if description:
                                desc_para = doc.add_paragraph()
                                desc_run = desc_para.add_run(description.strip())
                                desc_run.font.size = Pt(12)
                                desc_run.font.name = 'Calibri'
                                desc_para.paragraph_format.space_after = Pt(6)

                            else:
                                doc.add_paragraph("(Fallo la descripción)")
                        except ResourceExhausted as e:
                            print(f"Se excedio la cuota de API")
                            raise
                        except Exception as e:
                            print(f"Error generando descripción para {file}: {e}")
        
        conclusion_prompt = (
            f"Basándote en las visualizaciones y datos disponibles, redactá una conclusión profesional y clara que responda a la siguiente pregunta:\n"
            f"\"{question}\"\n\n"
            "Usá un tono técnico accesible, sin listas ni formato especial. Que parezca una conclusión escrita por un analista turístico."
        )

        try:
            conclusion_text = MetabaseDashboardExtract.get_image_description_from_gemini([], conclusion_prompt)
            if conclusion_text:
                conclusion_heading = doc.add_paragraph("Conclusión", style='Heading 1')
                conclusion_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                conclusion_heading.paragraph_format.space_after = Pt(12)
                    
                conclusion_para = doc.add_paragraph(conclusion_text.strip())
                conclusion_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            else:
                doc.add_paragraph("No se pudo generar una conclusión automática.")

        except Exception as e:
            doc.add_paragraph(f"(Error generando conclusión: {e})")


        safe_municipio = municipio.replace(" ", "_").lower()
        output_path = os.path.join(doc_dir, f"{safe_municipio}.docx")
        doc.save(output_path)

    


    def docx_to_pdf(self):
        input_path = "output/docs/teulada.docx"
        output_dir = "output/pdf"

        
        os.makedirs(output_dir, exist_ok=True)

        try:
            subprocess.run([
                "D:/Programas/LibreOffice/program/soffice.exe",       
                "--headless",           
                "--convert-to", "pdf",  
                "--outdir", output_dir,
                input_path
            ], check=True)

            print("todo ok")
        except subprocess.CalledProcessError as e:
            print(f"error libreoffice: {e}")

    
    def run(self):
        inicio = time.time()
        try:
            self.login()
            #municipios = ["teulada", "benissa", "calpe", "javea", "mijas", "moraira", "platja_de_aro"]
            municipios = ["teulada"]
            #tabs = ["100-vuts-y-casas-rurales", "102-hoteles%2C-hostales-y-campings", "103-datos-de-fuentes-oficiales"]
            tabs = ["100-vuts-y-casas-rurales"]

            logo_file_path = "logo_municipio.png"

            for municipio in municipios:
                for tab in tabs:
                    dashboard_url = f"https://analytics.peninsula.co/dashboard/32-{municipio}?a%25C3%25B1o=&fecha=&mes=&per%25C3%25ADodo=&poblaci%25C3%25B3n=teulada&poblaci%25C3%25B3n_%28consumo%29=calpe&poblaci%25C3%25B3n_2=benissa&tab={tab}&tipo_de_alojamiento=&tipo_de_establecimiento="
                    try:
                        self.capture_dashboard(dashboard_url, municipio, tab)
                    except Exception as e:
                        print(f"ERROR: {municipio} - {tab}: {e}")
                        break
            question = "¿Que tendencia se observa en la antelacion con la que los turistas reservan su alojamiento?"
            self.export_to_docx(question, logo_path=logo_file_path)
            self.docx_to_pdf()
            pdf_path = os.path.abspath("output/pdf/teulada.pdf")
            os.startfile(pdf_path)

        except ResourceExhausted as e:
            print(f"Se excedio la cuota de Gemini")
        except Exception as e:
            print(f"Error general: {e}")
        finally:
            self.driver.quit()
        fin = time.time()
        print(fin-inicio)

if __name__ == "__main__":
    email = os.getenv("METABASE_EMAIL")
    password = os.getenv("METABASE_PASSWORD")
    base_url = os.getenv("METABASE_BASE_URL")
    api_key = os.getenv("GEMINI_API_KEY")

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.0-flash')
    
    if not email or not password or not api_key or not base_url:
        raise ValueError("No estan las variables de entorno en el sistema")
        
    exporter = MetabaseDashboardExtract(
        email=email,
        password=password,
        base_url=base_url
    )
    exporter.run()